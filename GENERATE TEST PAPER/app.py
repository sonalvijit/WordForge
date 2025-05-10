from json import load
from docx import Document
from docx.shared import Pt

data:list = []

preposition:list = ["from","in","from","to","with","of","to","at",""]
dont_include = ["At a lost","At it again","Indifferent to"]
num:int = 1

with open("./phrases.json") as f:
     da_ = load(f)
     sr_ = sorted(da_, key=lambda x:x['phrase'].lower())
     for i in sr_:
          if num <= 10:
               if i['phrase'] in dont_include:
                    break
               ddd = i['phrase'].split(' ')
               if ddd[1] in preposition:
                    data.append({"phrase_questions":f"{ddd[0]} _______ "})
               num=num+1
          if num >=10:
               data.append({"phrase_questions":i['phrase']})
               num=num+1

def set_font_size(run, size):
    font = run.font
    font.size = Pt(size)

def create_docx_file(data, output_filename):
     doc = Document()
     doc.add_heading("TEST", 0)

     for entry in data:
          doc.add_heading(entry['phrase_questions'], level=1)
          p = doc.paragraphs[-1].runs[0]
          set_font_size(p, 14)

          p = doc.add_paragraph()
          run = p.add_run("Meaning: ")
          run.bold = True
          set_font_size(run, 14)

          p = doc.add_paragraph()
          run = p.add_run("Write sentence on your own words: ")
          # run.bold = True
          set_font_size(run, 12)

          for i in range(2):
               doc.add_paragraph(f" ", style='List Bullet').runs[0].font.size = Pt(12)

     doc.save(output_filename)


create_docx_file(data, "17032025.docx")