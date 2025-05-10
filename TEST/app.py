import json
from docx import Document
from docx.shared import Pt

# Function to read JSON data from a file
def load_json_data(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return json.load(file)

# Function to set font size for a given run
def set_font_size(run, size):
    font = run.font
    font.size = Pt(size)

# Function to create Word file from the data
def create_word_file(data, output_filename):
    doc = Document()
    doc.add_heading('Phrases and Their Meanings', 0)

    # Sort the data based on the 'phrase' field in alphabetical order
    sorted_data = sorted(data, key=lambda x: x['phrase'].lower())

    # Loop through each entry in the sorted data
    for entry in sorted_data:
        # Add phrase with default font size 14
        doc.add_heading(entry["phrase"], level=1)
        p = doc.paragraphs[-1].runs[0]
        set_font_size(p, 14)

        # Add meaning (with bold 'Meaning:' and font size 14)
        p = doc.add_paragraph()
        run = p.add_run("Meaning: ")
        run.bold = True
        set_font_size(run, 14)
        run = p.add_run(entry['meaning'])
        set_font_size(run, 12)

        # Add phrase type (with bold 'Phrase Type:' and font size 14)
        p = doc.add_paragraph()
        run = p.add_run("Phrase Type: ")
        run.bold = True
        set_font_size(run, 14)
        run = p.add_run(entry['phrase_type'])
        set_font_size(run, 12)

        # Add examples (with bold 'Examples:' and font size 14 for the key)
        p = doc.add_paragraph()
        run = p.add_run("Examples:")
        run.bold = True
        set_font_size(run, 14)
        for example in entry["example"]:
            doc.add_paragraph(f"- {example}", style='List Bullet').runs[0].font.size = Pt(12)

        # Add note if exists (with bold 'Note:' and font size 14 for the key)
        if entry["note"]:
            p = doc.add_paragraph()
            run = p.add_run("Note: ")
            run.bold = True
            set_font_size(run, 14)
            run = p.add_run(entry['note'])
            set_font_size(run, 12)
        
        # Add a separator between entries
        doc.add_paragraph("\n" + "-"*50 + "\n")
    
    # Save the document
    doc.save(output_filename)

# Main function to load the JSON file and create the Word file
def main():
    # Specify the path to your JSON file
    json_file_path = 'phrases.json'  # Replace with your actual file path
    data = load_json_data(json_file_path)
    
    # Call the function to create the Word file
    create_word_file(data, 'phrases.docx')

# Run the program
if __name__ == '__main__':
    main()
